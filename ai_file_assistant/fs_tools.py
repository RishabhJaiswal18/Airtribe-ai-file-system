# """
# fs_tools.py

# Core File System Tools for AI File Assistant

# This module provides structured tools for working with resume files.

# Tools Implemented:

# 1. read_file(filepath)
#    - Reads TXT, PDF, DOCX files
#    - Extracts text content
#    - Returns metadata

# 2. list_files(directory, extension=None)
#    - Lists files in directory
#    - Optional extension filtering
#    - Returns metadata

# 3. write_file(filepath, content)
#    - Writes TXT or DOCX files
#    - Creates directories if needed

# 4. search_in_file(filepath, keyword)
#    - Searches keyword in file
#    - Case insensitive search
#    - Returns context matches

# 5. search_in_directory(directory, keyword)
#    - Searches keyword across files
# """
import os
from datetime import datetime
import PyPDF2
import docx


# ===============================
# READ FILE TOOL
# ===============================


def read_file(filepath: str):

    try:

        if not os.path.exists(filepath):
            return {"status": "error", "message": "File not found"}

        ext = os.path.splitext(filepath)[1].lower()

        content = ""

        # TXT
        if ext == ".txt":

            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()

        # PDF
        elif ext == ".pdf":

            with open(filepath, "rb") as f:

                reader = PyPDF2.PdfReader(f)

                for page in reader.pages:
                    content += page.extract_text() or ""

        # DOCX
        elif ext == ".docx":

            document = docx.Document(filepath)

            for para in document.paragraphs:
                content += para.text + "\n"

        else:

            return {"status": "error", "message": "Unsupported file type"}

        metadata = {
            "filename": os.path.basename(filepath),
            "size_bytes": os.path.getsize(filepath),
            "modified_date": datetime.fromtimestamp(
                os.path.getmtime(filepath)
            ).isoformat(),
        }

        # Clean formatting
        content = " ".join(content.split())

        return {"status": "success", "content": content, "metadata": metadata}

    except Exception as e:

        return {"status": "error", "message": str(e)}


# ===============================
# LIST FILE TOOL
# ===============================


def list_files(directory: str, extension: str = None):

    try:

        if not os.path.exists(directory):
            return []

        files = []

        for file in os.listdir(directory):

            full = os.path.join(directory, file)

            if os.path.isfile(full):

                if extension:

                    if not extension.startswith("."):
                        extension = "." + extension

                    if not file.lower().endswith(extension.lower()):
                        continue

                files.append(
                    {
                        "name": file,
                        "size_bytes": os.path.getsize(full),
                        "modified_date": datetime.fromtimestamp(
                            os.path.getmtime(full)
                        ).isoformat(),
                    }
                )

        return files

    except Exception:

        return []


# ===============================
# WRITE FILE TOOL
# ===============================


def write_file(filepath: str, content: str):

    try:

        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        ext = os.path.splitext(filepath)[1].lower()

        # TXT
        if ext == ".txt":

            with open(filepath, "w", encoding="utf-8") as f:

                f.write(content)

        # DOCX
        elif ext == ".docx":

            document = docx.Document()

            document.add_paragraph(content)

            document.save(filepath)

        else:

            return {
                "status": "error",
                "message": "Only TXT and DOCX supported for writing",
            }

        return {"status": "success", "message": "File written successfully"}

    except Exception as e:

        return {"status": "error", "message": str(e)}


# ===============================
# SEARCH FILE TOOL
# ===============================


def search_in_file(filepath: str, keyword: str):

    try:

        data = read_file(filepath)

        if data["status"] != "success":
            return data

        text = data["content"]

        keyword_lower = keyword.lower()

        text_lower = text.lower()

        matches = []

        index = 0

        while True:

            index = text_lower.find(keyword_lower, index)

            if index == -1:
                break

            start = max(index - 40, 0)

            end = min(index + 40, len(text))

            # matches.append(text[start:end])
            matches.append(text[start:end].replace("\n", " ").strip())

            index += len(keyword_lower)

        return {
            "status": "success",
            "keyword": keyword,
            "matches": matches,
            "count": len(matches),
        }

    except Exception as e:

        return {"status": "error", "message": str(e)}


# ===============================
# SEARCH DIRECTORY TOOL
# ===============================


def search_in_directory(directory: str, keyword: str):

    try:

        results = []

        for file in os.listdir(directory):

            full = os.path.join(directory, file)

            if os.path.isfile(full):

                r = search_in_file(full, keyword)

                if r["status"] == "success" and r["count"] > 0:

                    results.append({"file": file, "matches": r["matches"]})

        return results

    except Exception as e:

        return {"status": "error", "message": str(e)}
